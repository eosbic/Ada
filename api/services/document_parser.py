"""
Document Parser — Parsea PDF, TXT, DOCX y extrae texto.
Se integra con el upload endpoint para análisis de cualquier documento.
"""

import os
import io
from typing import Tuple


def parse_document(file_bytes: bytes, file_name: str) -> Tuple[str, dict]:
    """
    Parsea un documento y extrae texto + metadata.
    
    Retorna: (texto_extraido, metadata)
    Soporta: .pdf, .txt, .md, .markdown, .doc, .docx
    """
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext == "txt":
        return _parse_txt(file_bytes, file_name)
    elif ext in ("md", "markdown"):
        return _parse_markdown(file_bytes, file_name)
    elif ext == "pdf":
        return _parse_pdf(file_bytes, file_name)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_bytes, file_name)
    else:
        return "", {"error": f"Formato .{ext} no soportado"}


def _parse_txt(file_bytes: bytes, file_name: str) -> Tuple[str, dict]:
    """Parsea archivo de texto plano."""
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        lines = text.strip().split("\n")
        return text, {
            "type": "txt",
            "file_name": file_name,
            "lines": len(lines),
            "chars": len(text),
            "words": len(text.split()),
        }
    except Exception as e:
        return "", {"error": f"Error leyendo TXT: {str(e)}"}


def _parse_markdown(file_bytes: bytes, file_name: str) -> Tuple[str, dict]:
    """Parsea Markdown como texto plano para analisis semantico."""
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        lines = text.strip().split("\n")
        return text, {
            "type": "markdown",
            "file_name": file_name,
            "lines": len(lines),
            "chars": len(text),
            "words": len(text.split()),
        }
    except Exception as e:
        return "", {"error": f"Error leyendo Markdown: {str(e)}"}


def _parse_pdf(file_bytes: bytes, file_name: str) -> Tuple[str, dict]:
    """Parsea PDF y extrae texto."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages_text.append(f"--- Página {page_num + 1} ---\n{text}")

        full_text = "\n\n".join(pages_text)

        metadata = {
            "type": "pdf",
            "file_name": file_name,
            "pages": len(doc),
            "chars": len(full_text),
            "words": len(full_text.split()),
        }

        doc.close()
        return full_text, metadata

    except ImportError:
        # Fallback sin PyMuPDF
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(io.BytesIO(file_bytes))
            return text, {
                "type": "pdf",
                "file_name": file_name,
                "chars": len(text),
                "words": len(text.split()),
            }
        except ImportError:
            return "", {"error": "No hay librería PDF instalada. Instalar: pip install PyMuPDF"}
        except Exception as e:
            return "", {"error": f"Error leyendo PDF: {str(e)}"}
    except Exception as e:
        return "", {"error": f"Error leyendo PDF: {str(e)}"}


def _parse_docx(file_bytes: bytes, file_name: str) -> Tuple[str, dict]:
    """Parsea DOCX/DOC y extrae texto."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extraer tablas también
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip().replace("|", "").strip():
                    tables_text.append(row_text)

        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n--- Tablas ---\n" + "\n".join(tables_text)

        return full_text, {
            "type": "docx",
            "file_name": file_name,
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables),
            "chars": len(full_text),
            "words": len(full_text.split()),
        }

    except ImportError:
        return "", {"error": "python-docx no instalado. Instalar: pip install python-docx"}
    except Exception as e:
        return "", {"error": f"Error leyendo DOCX: {str(e)}"}
